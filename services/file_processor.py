"""
File Processing Service — Extract text from documents, images, and markup files.
──────────────────────────────────────────────────────────────────────────────
Handles: .pdf, .docx, .xlsx, .xls, .txt, .csv, .json,
         .jpg, .jpeg, .png, .webp, .gif, .bmp, .tiff, .tif,
         .xml, .html, .htm, .svg, .md, .yaml, .yml, .log, .ini, .cfg
Produces: List[Chunk] ready for embedding.

Image files are processed via OpenAI Vision API (gpt-4o-mini) which
generates a detailed text description of the image content. This
description is then chunked and embedded like any text.
"""

import base64
import io
import json
import csv
import logging
import os
import re
from dataclasses import dataclass, field
from typing import List, Optional

import tiktoken
import chardet

logger = logging.getLogger("docqa.file_processor")

# tiktoken encoder for chunk size measurement
_enc = tiktoken.get_encoding("cl100k_base")

# Image extensions that use the vision extraction path
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tiff", ".tif"}

# MIME type mapping for base64 data URIs
_MIME_MAP = {
    ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".png": "image/png", ".webp": "image/webp",
    ".gif": "image/gif", ".bmp": "image/bmp",
    ".tiff": "image/tiff", ".tif": "image/tiff",
}


# ── Data Structures ───────────────────────────────────────────────────────────

@dataclass
class Chunk:
    text: str
    file_name: str
    file_type: str
    chunk_index: int
    page_number: Optional[int] = None
    sheet_name: Optional[str] = None
    char_start: int = 0
    char_end: int = 0
    token_count: int = 0


@dataclass
class ProcessedFile:
    file_name: str
    file_type: str
    size_bytes: int
    chunks: List[Chunk] = field(default_factory=list)
    status: str = "processed"
    error: Optional[str] = None


# ── Text Extractors ──────────────────────────────────────────────────────────

def extract_pdf(content: bytes, file_name: str) -> List[dict]:
    """Extract text from PDF, page by page. Returns [{text, page_number}]."""
    pages = []
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append({"text": text, "page_number": i})
    except Exception as e:
        logger.warning(f"pdfplumber failed for {file_name}, trying PyPDF2: {e}")
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append({"text": text, "page_number": i})
        except Exception as e2:
            logger.error(f"Both PDF extractors failed for {file_name}: {e2}")
            raise ValueError(f"Cannot extract text from PDF: {e2}")

    # If text extraction yielded very little (common for construction drawings/floor plans),
    # fall back to vision-based extraction for the entire PDF
    total_text = sum(len(p.get("text", "")) for p in pages)
    if total_text < 200:
        logger.info(f"PDF {file_name} has minimal text ({total_text} chars), using vision fallback")
        vision_pages = _extract_pdf_via_vision(content, file_name)
        if vision_pages:
            return vision_pages

    return pages


def _extract_pdf_via_vision(content: bytes, file_name: str) -> List[dict]:
    """Render PDF pages as images and extract content via OpenAI Vision API."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF (fitz) not installed — cannot use vision fallback for PDF")
        # Fall back to base64 encoding of first page as generic image
        return _extract_pdf_vision_simple(content, file_name)

    pages = []
    doc = fitz.open(stream=content, filetype="pdf")
    for page_num in range(min(len(doc), 20)):  # Cap at 20 pages
        page = doc[page_num]
        # Render page at 150 DPI (good balance of quality vs size)
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")
        description = _describe_image_sync(img_bytes, ".png", file_name, page_num + 1)
        if description:
            pages.append({"text": description, "page_number": page_num + 1})
    doc.close()
    return pages


def _extract_pdf_vision_simple(content: bytes, file_name: str) -> List[dict]:
    """Simple PDF vision fallback — send raw PDF bytes encoded as images if possible."""
    # For PDFs without PyMuPDF, try describing the file generically
    # We encode the first portion as a hint
    b64 = base64.b64encode(content[:5_000_000]).decode("ascii")  # Cap at 5MB
    description = _call_vision_api(
        b64, "application/pdf", file_name,
        "This is a construction document PDF. Describe ALL visible content: "
        "room labels, dimensions, annotations, notes, title block, drawing numbers, "
        "revision info, and any text visible on the drawing."
    )
    if description:
        return [{"text": description, "page_number": 1}]
    return []


def extract_docx(content: bytes, file_name: str) -> List[dict]:
    """Extract text from Word document. Returns [{text}]."""
    from docx import Document
    doc = Document(io.BytesIO(content))

    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    if not full_text.strip():
        raise ValueError(f"No text content found in {file_name}")
    return [{"text": full_text}]


def extract_excel(content: bytes, file_name: str) -> List[dict]:
    """Extract text from Excel, sheet by sheet. Returns [{text, sheet_name}]."""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)

    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        headers = None

        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            values = [str(v).strip() if v is not None else "" for v in row]
            if not any(values):
                continue

            if headers is None:
                headers = values
                rows_text.append(" | ".join(headers))
            else:
                # Label each cell with its header
                labeled = []
                for i, val in enumerate(values):
                    if val and i < len(headers) and headers[i]:
                        labeled.append(f"{headers[i]}: {val}")
                    elif val:
                        labeled.append(val)
                if labeled:
                    rows_text.append("; ".join(labeled))

        if rows_text:
            sheets.append({
                "text": "\n".join(rows_text),
                "sheet_name": sheet_name,
            })

    wb.close()
    if not sheets:
        raise ValueError(f"No data found in {file_name}")
    return sheets


def extract_text(content: bytes, file_name: str) -> List[dict]:
    """Extract text from plain text files (.txt, .csv, .json, .log, .ini, .cfg, .md, .yaml, .yml)."""
    # Detect encoding
    detected = chardet.detect(content)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    try:
        text = content.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        text = content.decode("utf-8", errors="replace")

    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else "txt"

    if ext == "csv":
        # Parse CSV into labeled rows
        reader = csv.DictReader(io.StringIO(text))
        rows = []
        for row in reader:
            labeled = "; ".join(f"{k}: {v}" for k, v in row.items() if v and v.strip())
            if labeled:
                rows.append(labeled)
        if rows:
            text = "\n".join(rows)

    elif ext == "json":
        try:
            data = json.loads(text)
            text = json.dumps(data, indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            pass  # Keep raw text

    if not text.strip():
        raise ValueError(f"No text content found in {file_name}")
    return [{"text": text}]


def extract_markup(content: bytes, file_name: str) -> List[dict]:
    """Extract text from XML, HTML, HTM, SVG files by stripping tags."""
    detected = chardet.detect(content)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    try:
        raw = content.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        raw = content.decode("utf-8", errors="replace")

    # Strip HTML/XML tags
    text = re.sub(r'<script[^>]*>.*?</script>', ' ', raw, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', ' ', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', ' ', text)
    # Decode common HTML entities
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    text = text.replace('&quot;', '"').replace('&nbsp;', ' ').replace('&#39;', "'")
    # Collapse whitespace
    text = re.sub(r'\s+', ' ', text).strip()

    if not text:
        raise ValueError(f"No text content found in {file_name}")
    return [{"text": text}]


# ── Image Extraction (Vision API) ─────────────────────────────────────────────

def _describe_image_sync(
    img_bytes: bytes,
    ext: str,
    file_name: str,
    page_number: Optional[int] = None,
) -> str:
    """Send an image to OpenAI Vision API and get a text description."""
    b64 = base64.b64encode(img_bytes).decode("ascii")
    mime = _MIME_MAP.get(ext, "image/png")
    page_hint = f" (page {page_number})" if page_number else ""
    prompt = (
        f"Analyze this construction/engineering document image '{file_name}'{page_hint} in detail. "
        "Extract and describe ALL visible content including:\n"
        "- Title block information (drawing number, project name, revision, date, scale)\n"
        "- Room names, labels, and numbers\n"
        "- Dimensions and measurements\n"
        "- Annotations, notes, and callouts\n"
        "- Material specifications and symbols\n"
        "- Equipment, fixtures, and their locations\n"
        "- Any legends, schedules, or tables\n"
        "- General layout description\n"
        "Be as thorough as possible. List every text element visible."
    )
    return _call_vision_api(b64, mime, file_name, prompt)


def _call_vision_api(b64_data: str, mime_type: str, file_name: str, prompt: str) -> str:
    """Call OpenAI Vision API synchronously. Returns description text."""
    try:
        from openai import OpenAI
        api_key = os.getenv("OPENAI_API_KEY", "")
        if not api_key:
            logger.error("OPENAI_API_KEY not set — cannot use vision extraction")
            return ""

        client = OpenAI(api_key=api_key)
        model = os.getenv("OPENAI_VISION_MODEL", "gpt-4o")

        response = client.chat.completions.create(
            model=model,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {
                        "url": f"data:{mime_type};base64,{b64_data}",
                        "detail": "high",
                    }},
                ],
            }],
            max_tokens=4096,
            temperature=0.1,
        )
        description = response.choices[0].message.content or ""
        logger.info(f"Vision API extracted {len(description)} chars from {file_name}")
        return description

    except Exception as e:
        logger.error(f"Vision API call failed for {file_name}: {e}")
        return ""


def extract_image(content: bytes, file_name: str) -> List[dict]:
    """Extract text from image files via OpenAI Vision API."""
    ext = "." + file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ".png"
    description = _describe_image_sync(content, ext, file_name)
    if not description:
        raise ValueError(
            f"Could not extract content from image {file_name}. "
            "Ensure OPENAI_API_KEY is configured."
        )
    return [{"text": description}]


# ── Parser Registry ──────────────────────────────────────────────────────────

PARSERS = {
    # Documents
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_excel,
    ".xls": extract_excel,
    # Plain text
    ".txt": extract_text,
    ".csv": extract_text,
    ".json": extract_text,
    ".md": extract_text,
    ".yaml": extract_text,
    ".yml": extract_text,
    ".log": extract_text,
    ".ini": extract_text,
    ".cfg": extract_text,
    # Markup
    ".xml": extract_markup,
    ".html": extract_markup,
    ".htm": extract_markup,
    ".svg": extract_markup,
    # Images (via Vision API)
    ".jpg": extract_image,
    ".jpeg": extract_image,
    ".png": extract_image,
    ".webp": extract_image,
    ".gif": extract_image,
    ".bmp": extract_image,
    ".tiff": extract_image,
    ".tif": extract_image,
}


# ── Chunking ─────────────────────────────────────────────────────────────────

def _split_into_sentences(text: str) -> List[str]:
    """Split text into sentences at .!? boundaries."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s for s in sentences if s.strip()]


def chunk_text(
    text: str,
    chunk_size: int = 512,
    overlap: int = 64,
) -> List[str]:
    """
    Split text into token-measured chunks with overlap.
    Uses sentence boundaries for cleaner splits.
    """
    sentences = _split_into_sentences(text)
    if not sentences:
        return []

    chunks = []
    current_tokens = []
    current_count = 0

    for sentence in sentences:
        sent_tokens = _enc.encode(sentence)
        sent_len = len(sent_tokens)

        # If single sentence exceeds chunk_size, split it by tokens
        if sent_len > chunk_size:
            # Flush current buffer
            if current_tokens:
                chunks.append(_enc.decode(current_tokens))
                # Keep overlap
                if overlap > 0 and len(current_tokens) > overlap:
                    current_tokens = current_tokens[-overlap:]
                    current_count = len(current_tokens)
                else:
                    current_tokens = []
                    current_count = 0

            # Split long sentence into fixed-size pieces
            for i in range(0, sent_len, chunk_size - overlap):
                piece = sent_tokens[i:i + chunk_size]
                chunks.append(_enc.decode(piece))
            current_tokens = []
            current_count = 0
            continue

        if current_count + sent_len > chunk_size:
            # Flush
            chunks.append(_enc.decode(current_tokens))
            # Keep overlap tokens from end
            if overlap > 0 and len(current_tokens) > overlap:
                current_tokens = current_tokens[-overlap:]
                current_count = len(current_tokens)
            else:
                current_tokens = []
                current_count = 0

        current_tokens.extend(sent_tokens)
        current_count += sent_len

    # Flush remainder
    if current_tokens:
        chunks.append(_enc.decode(current_tokens))

    return chunks


# ── Main Processor ───────────────────────────────────────────────────────────

class FileProcessor:
    """Process uploaded files into chunks ready for embedding."""

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def process(self, file_name: str, content: bytes) -> ProcessedFile:
        """Process a single file into chunks."""
        ext = ""
        if "." in file_name:
            ext = "." + file_name.rsplit(".", 1)[-1].lower()

        result = ProcessedFile(
            file_name=file_name,
            file_type=ext,
            size_bytes=len(content),
        )

        parser = PARSERS.get(ext)
        if parser is None:
            result.status = "unsupported"
            result.error = f"Unsupported file type: {ext}"
            return result

        try:
            sections = parser(content, file_name)
        except Exception as e:
            result.status = "failed"
            result.error = str(e)
            logger.error(f"Failed to extract text from {file_name}: {e}")
            return result

        chunk_index = 0
        for section in sections:
            text = section.get("text", "")
            page_number = section.get("page_number")
            sheet_name = section.get("sheet_name")

            text_chunks = chunk_text(text, self.chunk_size, self.chunk_overlap)

            for chunk_text_str in text_chunks:
                token_count = len(_enc.encode(chunk_text_str))
                chunk = Chunk(
                    text=chunk_text_str,
                    file_name=file_name,
                    file_type=ext,
                    chunk_index=chunk_index,
                    page_number=page_number,
                    sheet_name=sheet_name,
                    token_count=token_count,
                )
                result.chunks.append(chunk)
                chunk_index += 1

        if not result.chunks:
            result.status = "failed"
            result.error = "No text content extracted"

        logger.info(f"Processed {file_name}: {len(result.chunks)} chunks")
        return result
